#!/usr/bin/env python3
"""Generate independent PDF, ePub, and DOCX fixtures for document translation tests."""

from __future__ import annotations

import argparse
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas


def set_run_font(run, name: str, size: float, color: str = "202533", bold: bool = False) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def generate_docx(output: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, "2E74B5", 18, 10),
        "Heading 2": (13, "2E74B5", 14, 7),
        "Heading 3": (12, "1F4D78", 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run("FLUENTREAD · DOCUMENT TRANSLATION EXAMPLE"), "Calibri", 8.5, "7A8294", True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    set_run_font(footer.add_run("Local browser fixture"), "Calibri", 8.5, "7A8294")

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(6)
    set_run_font(kicker.add_run("REFERENCE GUIDE"), "Calibri", 10, "E83B6B", True)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    set_run_font(title.add_run("Document Translation Example"), "Calibri", 26, "202533", True)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(18)
    set_run_font(subtitle.add_run("A realistic DOCX fixture for bilingual reading regression"), "Calibri", 13, "5F687C")

    document.add_heading("Why local document translation matters", level=1)
    document.add_paragraph(
        "A useful document translator should preserve the source file while presenting each translated paragraph directly beneath its original text."
    )
    paragraph = document.add_paragraph()
    set_run_font(paragraph.add_run("Privacy first. "), "Calibri", 11, "202533", True)
    set_run_font(
        paragraph.add_run("The file is parsed inside the browser, and only extracted text is sent to the translation service selected by the reader."),
        "Calibri",
        11,
    )

    document.add_heading("What the regression should protect", level=1)
    for label, detail in (
        ("Structure", "Headings, paragraph order, headers, and footers remain valid after export."),
        ("Reading flow", "The original paragraph appears first and the editable translation follows naturally."),
        ("Format", "The downloaded result remains a standards-compliant DOCX that opens in Word and LibreOffice."),
    ):
        paragraph = document.add_paragraph()
        set_run_font(paragraph.add_run(f"{label}: "), "Calibri", 11, "2E74B5", True)
        set_run_font(paragraph.add_run(detail), "Calibri", 11)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def generate_pdf(output: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    page_width, page_height = letter
    pdf = canvas.Canvas(str(output), pagesize=letter)
    pdf.setTitle("Document Translation Example")
    pdf.setAuthor("FluentRead")

    def wrapped_lines(text: str, font: str, size: float, width: float) -> list[str]:
        words = text.split()
        lines: list[str] = []
        current = ""
        for word in words:
            candidate = f"{current} {word}".strip()
            if current and stringWidth(candidate, font, size) > width:
                lines.append(current)
                current = word
            else:
                current = candidate
        if current:
            lines.append(current)
        return lines

    def draw_wrapped(text: str, x: float, y: float, width: float, *, font: str = "Times-Roman", size: float = 9.2, leading: float = 12.3) -> float:
        pdf.setFont(font, size)
        pdf.setFillColor(HexColor("#202533"))
        for line in wrapped_lines(text, font, size, width):
            pdf.drawString(x, y, line)
            y -= leading
        return y

    def draw_heading(text: str, x: float, y: float) -> float:
        pdf.setFont("Helvetica-Bold", 11)
        pdf.setFillColor(HexColor("#D93463"))
        pdf.drawString(x, y, text)
        return y - 17

    def draw_footer(page_number: int) -> None:
        pdf.setStrokeColor(HexColor("#DDE2EC"))
        pdf.line(48, 39, page_width - 48, 39)
        pdf.setFont("Helvetica", 7.5)
        pdf.setFillColor(HexColor("#7A8294"))
        pdf.drawString(48, 25, "FLUENTREAD DOCUMENT LAYOUT REGRESSION")
        pdf.drawRightString(page_width - 48, 25, str(page_number))

    # 第 1 页刻意组合通栏标题、双栏正文和图示，用于捕获 PDF 翻译被压平成
    # 提取字符串列表、而没有写回原页面文本框的回归。
    pdf.setFont("Helvetica-Bold", 22)
    pdf.setFillColor(HexColor("#202533"))
    pdf.drawCentredString(page_width / 2, 742, "Document Translation Example")
    pdf.setFont("Helvetica", 9)
    pdf.setFillColor(HexColor("#647086"))
    pdf.drawCentredString(page_width / 2, 724, "Layout-preserving bilingual PDF regression · FluentRead Research")
    pdf.setFillColor(HexColor("#F7F8FB"))
    pdf.roundRect(48, 652, page_width - 96, 51, 7, fill=1, stroke=0)
    pdf.setFillColor(HexColor("#202533"))
    pdf.setFont("Helvetica-Bold", 8.5)
    pdf.drawString(60, 688, "ABSTRACT")
    draw_wrapped(
        "A document translator should preserve columns, figures, captions, and reading order while replacing text inside the corresponding page regions.",
        60, 674, page_width - 120, font="Helvetica", size=8.4, leading=10.4,
    )

    left_x, right_x, column_width = 48, 318, 246
    left_y = draw_heading("1  PAGE-AWARE TRANSLATION", left_x, 626)
    left_y = draw_wrapped(
        "FluentRead parses the PDF text layer into layout blocks with page coordinates. Each translated block is fitted back into its original region instead of being shown as a detached transcript.",
        left_x, left_y, column_width,
    ) - 9
    left_y = draw_heading("2  READING EXPERIENCE", left_x, left_y)
    left_y = draw_wrapped(
        "The reader presents the untouched source page and the translated page side by side. Pagination, multi-column flow, graphics, and captions remain visible for visual comparison.",
        left_x, left_y, column_width,
    ) - 9
    left_y = draw_heading("3  EXPORT CONTRACT", left_x, left_y)
    draw_wrapped(
        "A bilingual download contains one wide page for every source page: the original is placed on the left and the translated layout is placed on the right.",
        left_x, left_y, column_width,
    )

    right_y = draw_heading("FIGURE 1  DOCUMENT PIPELINE", right_x, 626)
    box_y = right_y - 56
    box_width = 66
    for index, (label, color) in enumerate((("Parse", "#EAF0FF"), ("Translate", "#FFF0F4"), ("Render", "#E7F8F5"))):
        box_x = right_x + index * 84
        pdf.setFillColor(HexColor(color))
        pdf.setStrokeColor(HexColor("#C9D1E2"))
        pdf.roundRect(box_x, box_y, box_width, 38, 6, fill=1, stroke=1)
        pdf.setFillColor(HexColor("#202533"))
        pdf.setFont("Helvetica-Bold", 8)
        pdf.drawCentredString(box_x + box_width / 2, box_y + 15, label)
        if index < 2:
            pdf.setStrokeColor(HexColor("#9BA6BB"))
            pdf.line(box_x + box_width + 5, box_y + 19, box_x + box_width + 15, box_y + 19)
            pdf.line(box_x + box_width + 11, box_y + 22, box_x + box_width + 15, box_y + 19)
            pdf.line(box_x + box_width + 11, box_y + 16, box_x + box_width + 15, box_y + 19)
    right_y = box_y - 19
    pdf.setFont("Times-Italic", 8.2)
    pdf.setFillColor(HexColor("#5F687C"))
    pdf.drawString(right_x, right_y, "Text extraction is an internal step, not the final reading interface.")
    right_y -= 29
    right_y = draw_heading("4  VISUAL INVARIANTS", right_x, right_y)
    draw_wrapped(
        "Regression evidence must prove that the translated page retains the figure, column boundaries, title hierarchy, and page dimensions. A generic list of OCR-like text fragments does not satisfy this contract.",
        right_x, right_y, column_width,
    )
    draw_footer(1)
    pdf.showPage()

    # 第 2 页覆盖表格、图表、通栏说明和独立分栏，确保解析与导出定位都受到验证。
    pdf.setFont("Helvetica-Bold", 19)
    pdf.setFillColor(HexColor("#202533"))
    pdf.drawString(48, 742, "Regression Coverage")
    pdf.setFont("Helvetica", 9)
    pdf.setFillColor(HexColor("#647086"))
    pdf.drawString(48, 724, "Expected invariants for preview and downloadable PDF output")

    table_x, table_y = 48, 634
    widths = [155, 105, 248]
    row_height = 25
    headers = ["Capability", "Expected", "Evidence"]
    rows = [
        ("Page dimensions", "Preserved", "Source and translated pages share the same height"),
        ("Two-column flow", "Preserved", "Blocks remain in their original left or right column"),
        ("Figures and charts", "Preserved", "Non-text page graphics remain in place"),
    ]
    pdf.setFillColor(HexColor("#202533"))
    pdf.rect(table_x, table_y + row_height, sum(widths), row_height, fill=1, stroke=0)
    cursor_x = table_x
    pdf.setFont("Helvetica-Bold", 8)
    pdf.setFillColor(HexColor("#FFFFFF"))
    for header, width in zip(headers, widths):
        pdf.drawString(cursor_x + 8, table_y + 34, header)
        cursor_x += width
    for row_index, row in enumerate(rows):
        row_y = table_y - row_index * row_height
        pdf.setFillColor(HexColor("#F7F8FB") if row_index % 2 == 0 else HexColor("#FFFFFF"))
        pdf.setStrokeColor(HexColor("#DDE2EC"))
        pdf.rect(table_x, row_y, sum(widths), row_height, fill=1, stroke=1)
        cursor_x = table_x
        pdf.setFillColor(HexColor("#202533"))
        pdf.setFont("Helvetica", 7.8)
        for value, width in zip(row, widths):
            pdf.drawString(cursor_x + 8, row_y + 9, value)
            cursor_x += width

    left_y = draw_heading("5  MULTI-COLUMN CONTENT", left_x, 527)
    left_y = draw_wrapped(
        "Page boundaries and paragraph order survive parsing. The translated page uses the same canvas as the source page before text regions are replaced, which leaves lines, colored shapes, and charts untouched.",
        left_x, left_y, column_width,
    ) - 8
    left_y = draw_heading("6  FAILURE BOUNDARY", left_x, left_y)
    draw_wrapped(
        "Image-only scans without a usable text layer are reported clearly. This beta does not pretend that extracted OCR text alone is a translated PDF document.",
        left_x, left_y, column_width,
    )

    right_y = draw_heading("FIGURE 2  BLOCK RETENTION", right_x, 527)
    chart_x, chart_y, chart_w, chart_h = right_x + 12, 382, 215, 92
    pdf.setStrokeColor(HexColor("#AAB3C5"))
    pdf.line(chart_x, chart_y, chart_x, chart_y + chart_h)
    pdf.line(chart_x, chart_y, chart_x + chart_w, chart_y)
    for index, (label, value, color) in enumerate((("Text", .86, "#ED4775"), ("Layout", .96, "#5A7BEF"), ("Graphics", 1.0, "#2DB483"))):
        bar_x = chart_x + 25 + index * 62
        pdf.setFillColor(HexColor(color))
        pdf.rect(bar_x, chart_y, 29, chart_h * value, fill=1, stroke=0)
        pdf.setFillColor(HexColor("#5F687C"))
        pdf.setFont("Helvetica", 7)
        pdf.drawCentredString(bar_x + 14.5, chart_y - 11, label)
    pdf.setFont("Times-Italic", 8.1)
    pdf.setFillColor(HexColor("#5F687C"))
    pdf.drawString(right_x, 350, "Figure 2. Layout and non-text graphics remain anchored to the page.")
    draw_footer(2)
    pdf.save()


def generate_epub(output: Path) -> None:
    mimetype = "application/epub+zip"
    container = """<?xml version="1.0" encoding="UTF-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles><rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/></rootfiles>
</container>
"""
    content_opf = """<?xml version="1.0" encoding="UTF-8"?>
<package xmlns="http://www.idpf.org/2007/opf" version="3.0" unique-identifier="book-id">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
    <dc:identifier id="book-id">fluentread-document-example</dc:identifier>
    <dc:title>Document Translation Example</dc:title>
    <dc:language>en</dc:language>
  </metadata>
  <manifest>
    <item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>
    <item id="chapter-1" href="chapter-1.xhtml" media-type="application/xhtml+xml"/>
    <item id="chapter-2" href="chapter-2.xhtml" media-type="application/xhtml+xml"/>
  </manifest>
  <spine>
    <itemref idref="chapter-1"/>
    <itemref idref="chapter-2"/>
  </spine>
</package>
"""
    nav = """<?xml version="1.0" encoding="UTF-8"?>
<html xmlns="http://www.w3.org/1999/xhtml"><head><title>Contents</title></head>
<body><nav epub:type="toc" xmlns:epub="http://www.idpf.org/2007/ops"><ol>
<li><a href="chapter-1.xhtml">Fluent reading</a></li><li><a href="chapter-2.xhtml">Regression</a></li>
</ol></nav></body></html>
"""
    chapter_1 = """<?xml version="1.0" encoding="UTF-8"?>
<html xmlns="http://www.w3.org/1999/xhtml"><head><title>Fluent reading</title></head><body>
<h1>Fluent reading for local books</h1>
<p>An ePub translator should preserve chapter order, links, and the original source while placing each translation directly below its paragraph.</p>
<p>Readers can edit the translation before downloading a bilingual electronic book.</p>
</body></html>
"""
    chapter_2 = """<?xml version="1.0" encoding="UTF-8"?>
<html xmlns="http://www.w3.org/1999/xhtml"><head><title>Regression coverage</title></head><body>
<h1>Regression coverage</h1>
<p>The exported archive must remain a valid ePub package with an uncompressed mimetype entry and readable XHTML chapters.</p>
<p><a href="chapter-1.xhtml">Return to the first chapter</a> without changing the link target.</p>
</body></html>
"""

    output.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(output, "w") as archive:
        archive.writestr("mimetype", mimetype, compress_type=zipfile.ZIP_STORED)
        archive.writestr("META-INF/container.xml", container, compress_type=zipfile.ZIP_DEFLATED)
        archive.writestr("OEBPS/content.opf", content_opf, compress_type=zipfile.ZIP_DEFLATED)
        archive.writestr("OEBPS/nav.xhtml", nav, compress_type=zipfile.ZIP_DEFLATED)
        archive.writestr("OEBPS/chapter-1.xhtml", chapter_1, compress_type=zipfile.ZIP_DEFLATED)
        archive.writestr("OEBPS/chapter-2.xhtml", chapter_2, compress_type=zipfile.ZIP_DEFLATED)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", type=Path, required=True)
    args = parser.parse_args()
    generate_pdf(args.output_dir / "sample.pdf")
    generate_epub(args.output_dir / "sample.epub")
    generate_docx(args.output_dir / "sample.docx")


if __name__ == "__main__":
    main()
